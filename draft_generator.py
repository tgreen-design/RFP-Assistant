"""
draft_generator.py
────────────────────────────────────────────────────────────────
Generates a formatted Word (.docx) RFP response draft.

AUTOPILOT sections are pre-filled from config.py boilerplate.
NEEDS_INPUT sections are flagged with yellow highlight boxes.
LEGAL_FLAG sections are flagged with red highlight boxes.

Returns: path to the generated .docx file
"""

import sys, json
from pathlib import Path
from datetime import datetime

sys.path.insert(0, str(Path(__file__).parent))
import config

OUTPUT_DIR = Path(__file__).parent / "rfp_output"
OUTPUT_DIR.mkdir(exist_ok=True)

# ─────────────────────────────────────────────────────────────────────────────
#  BUILD THE DRAFT USING python-docx
# ─────────────────────────────────────────────────────────────────────────────
def generate_draft(parsed_rfp: dict, tasks: list, go_no_go: str, questions: str) -> Path:
    """
    Generate the full RFP response draft as a .docx file.
    Returns the path to the saved file.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print("ERROR: python-docx not installed. Run: pip install python-docx")
        return None

    company   = parsed_rfp.get("company_name", "Unknown Company")
    deadline  = parsed_rfp.get("submission_deadline", "Unknown")
    industry  = parsed_rfp.get("industry", "Unknown")
    sections  = parsed_rfp.get("sections", [])

    # Map section names to task categories
    task_map = {}
    for t in tasks:
        task_map[t["section_name"].lower()] = t

    # ── Helper: find task category for a section ─────────────────────────────
    def get_category(section_name: str) -> dict:
        """Find the task entry for a given section name (fuzzy match)."""
        lower = section_name.lower()
        for key, task in task_map.items():
            if lower in key or key in lower:
                return task
        # Default: needs input
        return {"category": "NEEDS_INPUT", "owner": "Ryan Parrish",
                "priority": "medium", "notes": ""}

    # ── Helper: add colored section banner ───────────────────────────────────
    def add_banner(doc, text, category):
        colors = {
            "AUTOPILOT":   (0xE2, 0xEF, 0xDA),  # light green
            "NEEDS_INPUT": (0xFF, 0xF2, 0xCC),   # light yellow
            "LEGAL_FLAG":  (0xFC, 0xE4, 0xD6),   # light red/orange
        }
        text_colors = {
            "AUTOPILOT":   (0x37, 0x56, 0x23),
            "NEEDS_INPUT": (0x7F, 0x60, 0x00),
            "LEGAL_FLAG":  (0x84, 0x3C, 0x0C),
        }
        labels = {
            "AUTOPILOT":   "AUTO-FILLED",
            "NEEDS_INPUT": "NEEDS INPUT",
            "LEGAL_FLAG":  "LEGAL REVIEW REQUIRED",
        }

        bg = colors.get(category, colors["NEEDS_INPUT"])
        fg = text_colors.get(category, text_colors["NEEDS_INPUT"])
        label = labels.get(category, "")

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(f"  {label}  |  {text}")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(*fg)

        # Set background shading via XML
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "{:02X}{:02X}{:02X}".format(*bg))
        pPr.append(shd)
        return p

    # ── Helper: add a flag box ────────────────────────────────────────────────
    def add_flag_box(doc, category, owner, notes):
        colors = {
            "NEEDS_INPUT": (0xFF, 0xF2, 0xCC),
            "LEGAL_FLAG":  (0xFC, 0xE4, 0xD6),
        }
        instructions = {
            "NEEDS_INPUT": (
                f"[FILL IN — Owner: {owner}]\n"
                f"This section requires custom content. {notes if notes else ''}"
            ),
            "LEGAL_FLAG": (
                f"[LEGAL REVIEW REQUIRED — Owner: {owner}]\n"
                f"This section contains non-standard terms. Do not draft response until Legal has reviewed.\n"
                f"{notes if notes else ''}"
            ),
        }

        bg  = colors.get(category, colors["NEEDS_INPUT"])
        msg = instructions.get(category, "[FILL IN]")

        p = doc.add_paragraph()
        run = p.add_run(msg)
        run.font.size = Pt(10)
        run.italic    = True

        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "{:02X}{:02X}{:02X}".format(*bg))
        pPr.append(shd)

    # ── Helper: add autopilot content ────────────────────────────────────────
    def add_autopilot_content(doc, section_name: str, section_content: str):
        """Intelligently map section to boilerplate and add it."""
        lower = section_name.lower()
        content = None

        # Map section names to config.py boilerplate keys
        if any(w in lower for w in ["company info", "company overview", "about us", "vendor company"]):
            content = config.AUTOPILOT_CONTENT.get("company_overview_long")
        elif any(w in lower for w in ["contact", "primary contact"]):
            content = (f"{config.PRIMARY_CONTACT_NAME}\n"
                       f"{config.PRIMARY_CONTACT_TITLE}\n"
                       f"{config.PRIMARY_CONTACT_PHONE} | {config.PRIMARY_CONTACT_EMAIL}")
        elif any(w in lower for w in ["fulfillment", "logistics", "delivery", "shipping", "timing"]):
            content = config.AUTOPILOT_CONTENT.get("fulfillment_process")
        elif any(w in lower for w in ["sustain", "environmental", "esg"]):
            content = config.AUTOPILOT_CONTENT.get("sustainability")
        elif any(w in lower for w in ["security", "compliance", "soc", "hipaa", "audit"]):
            content = config.AUTOPILOT_CONTENT.get("security_and_compliance")
        elif any(w in lower for w in ["financial", "stability", "bankruptcy", "litigation"]):
            content = config.AUTOPILOT_CONTENT.get("financial_disclosures")
        elif any(w in lower for w in ["portal", "ordering", "reporting", "ecommerce", "platform"]):
            content = config.AUTOPILOT_CONTENT.get("portal_capabilities")
        elif any(w in lower for w in ["design", "creative", "artwork", "artist"]):
            content = config.AUTOPILOT_CONTENT.get("design_capabilities")
        elif any(w in lower for w in ["catalog", "product", "brand", "assortment"]):
            content = config.AUTOPILOT_CONTENT.get("product_catalog_reference")
        elif any(w in lower for w in ["payment", "terms", "net ", "invoic"]):
            content = config.AUTOPILOT_CONTENT.get("payment_terms_standard")
        elif any(w in lower for w in ["background", "employee", "conduct"]):
            content = config.AUTOPILOT_CONTENT.get("background_checks")
        elif any(w in lower for w in ["reference", "client list", "customer list"]):
            content = config.AUTOPILOT_CONTENT.get("client_confidentiality")
        elif any(w in lower for w in ["satisfaction", "guarantee", "damage", "defect", "penalty"]):
            content = config.AUTOPILOT_CONTENT.get("satisfaction_guarantee")

        if content:
            p = doc.add_paragraph(content)
            p.paragraph_format.space_after = Pt(6)
        else:
            # Fall back to a generic placeholder with the section content as context
            p = doc.add_paragraph(
                f"[Standard response applies. Review and customize if needed.]\n"
                f"Context from RFP: {section_content[:300]}..."
            )
            p.runs[0].italic = True

    # ─────────────────────────────────────────────────────────────────────────
    #  BUILD THE DOCUMENT
    # ─────────────────────────────────────────────────────────────────────────
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── COVER ──────────────────────────────────────────────────────────────
    title = doc.add_heading("CustomInk — RFP Response", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading(f"Prepared for: {company}", level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Submission Deadline: {deadline}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Prepared by: {config.PRIMARY_CONTACT_NAME}, {config.PRIMARY_CONTACT_TITLE}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    # ── KEY FACTS / GO NO GO ──────────────────────────────────────────────
    doc.add_heading("Deal Overview", level=2)

    # Extract verdict
    verdict = "YELLOW"
    for line in go_no_go.split("\n"):
        if line.startswith("RECOMMENDATION:"):
            verdict = line.replace("RECOMMENDATION:", "").strip()
            break

    labels = {"GREEN": "GO — Pursue", "YELLOW": "CONDITIONAL — Review required", "RED": "NO-GO — Do not pursue"}
    overview_p = doc.add_paragraph(f"Agent Recommendation: {verdict} — {labels.get(verdict, '')}")
    overview_p.runs[0].bold = True

    doc.add_paragraph(f"Industry: {industry}")
    doc.add_paragraph(f"Estimated Volume: {parsed_rfp.get('estimated_annual_volume', 'Unknown')}")
    doc.add_paragraph(f"Geographic Scope: {parsed_rfp.get('geographic_scope', 'Unknown')}")
    doc.add_paragraph("")

    # ── LEGAL FLAGS CALLOUT ───────────────────────────────────────────────
    legal_flags = parsed_rfp.get("legal_flags", [])
    if legal_flags:
        doc.add_heading("Legal Flags — Route to Legal Before Submitting", level=2)
        for flag in legal_flags:
            p = doc.add_paragraph(f"• {flag}", style="List Bullet")
        doc.add_paragraph("")

    # ── MISSING INFO ──────────────────────────────────────────────────────
    missing = parsed_rfp.get("missing_info", [])
    if missing:
        doc.add_heading("Information Gaps — Clarify with Prospect", level=2)
        for m in missing:
            doc.add_paragraph(f"• {m}", style="List Bullet")
        doc.add_paragraph("")

    doc.add_page_break()

    # ── RESPONSE SECTIONS ────────────────────────────────────────────────
    doc.add_heading("Response Sections", level=1)
    doc.add_paragraph(
        "Color coding: Green = auto-filled from standard boilerplate | "
        "Yellow = needs custom content from your team | "
        "Orange/Red = requires Legal review before completing"
    ).italic = True
    doc.add_paragraph("")

    for section in sections:
        name    = section.get("section_name", "Unnamed Section")
        content = section.get("section_content", "")
        task    = get_category(name)
        cat     = task.get("category", "NEEDS_INPUT")
        owner   = task.get("owner", "Ryan Parrish")
        notes   = task.get("notes", "")

        # Section heading + colored banner
        doc.add_heading(name, level=2)
        add_banner(doc, name, cat)

        if cat == "AUTOPILOT":
            add_autopilot_content(doc, name, content)
        else:
            # Show what the RFP is asking for, then the flag
            if content:
                context_p = doc.add_paragraph(f"RFP asks for: {content[:400]}")
                context_p.runs[0].italic = True
                context_p.runs[0].font.size = Pt(9)
            add_flag_box(doc, cat, owner, notes)

        doc.add_paragraph("")

    doc.add_page_break()

    # ── CLARIFICATION QUESTIONS ───────────────────────────────────────────
    doc.add_heading("Clarification Questions — Ready to Send to Prospect", level=1)
    doc.add_paragraph(
        "Review the internal routing note at the bottom before sending. "
        "Remove the [INTERNAL ROUTING NOTE] section before emailing to the prospect."
    ).italic = True
    doc.add_paragraph("")

    # Split into public + internal
    public_q   = questions
    internal_q = ""
    if "INTERNAL ROUTING NOTE" in questions:
        parts      = questions.split("INTERNAL ROUTING NOTE", 1)
        public_q   = parts[0].strip()
        internal_q = parts[1].strip()

    doc.add_paragraph(public_q)
    doc.add_paragraph("")

    if internal_q:
        doc.add_heading("INTERNAL ROUTING NOTE — Do Not Send to Prospect", level=2)
        doc.add_paragraph(internal_q)

    doc.add_page_break()

    # ── GO/NO-GO BRIEF ───────────────────────────────────────────────────
    doc.add_heading("Go/No-Go Brief (Full)", level=1)
    doc.add_paragraph(go_no_go)

    # ── SAVE ─────────────────────────────────────────────────────────────
    safe_name  = company.replace(" ", "_").replace("/", "-")[:30]
    date_str   = datetime.now().strftime("%Y%m%d_%H%M")
    filename   = OUTPUT_DIR / f"{safe_name}_{date_str}_DRAFT.docx"
    doc.save(filename)
    return filename


# ─────────────────────────────────────────────────────────────────────────────
#  STANDALONE TEST
# ─────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    # Quick smoke test with dummy data
    dummy_parsed = {
        "company_name": "Test Company",
        "industry": "Technology",
        "company_size": "enterprise",
        "submission_deadline": "2026-06-01",
        "estimated_annual_volume": "$500,000/year",
        "geographic_scope": "US only",
        "legal_flags": ["Unlimited indemnification clause", "IP assignment on all designs"],
        "missing_info": ["Current vendor name", "Number of employees using program"],
        "sections": [
            {"section_name": "Company Overview", "section_content": "Describe your company history and size."},
            {"section_name": "Pricing Proposal", "section_content": "Provide per-unit pricing by category."},
            {"section_name": "Contract Terms", "section_content": "Unlimited indemnification required."},
        ]
    }
    dummy_tasks = [
        {"section_name": "Company Overview", "category": "AUTOPILOT", "owner": "Ryan Parrish", "priority": "low", "notes": ""},
        {"section_name": "Pricing Proposal", "category": "NEEDS_INPUT", "owner": "Cody Perry", "priority": "high", "notes": "Build pricing matrix"},
        {"section_name": "Contract Terms", "category": "LEGAL_FLAG", "owner": "Legal Team", "priority": "high", "notes": "Unlimited indemnification — must counter"},
    ]
    path = generate_draft(dummy_parsed, dummy_tasks, "RECOMMENDATION: YELLOW\n\nDEAL SNAPSHOT\n- Company: Test\n", "1. What is X?\n2. What is Y?\n\nINTERNAL ROUTING NOTE\nRoute Q2 to Legal.")
    if path:
        print(f"Draft saved to: {path}")
    else:
        print("Draft generation failed.")
